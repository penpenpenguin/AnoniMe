import os
import json
import shutil
from datetime import datetime
import tempfile
from pathlib import Path
import sys  # 新增：供平台判斷
import subprocess  # 新增：啟動外部處理腳本
from urllib.parse import urlparse, unquote  # 新增：解析 file:// URL
import re  # 新增：解析 stdout 中的路徑
from zipfile import ZipFile, ZIP_DEFLATED

from PySide6.QtCore import QObject, Signal, Slot

try:
    from file_handlers.txt_handler import TextHandler
except ImportError:
    print("警告：無法導入 TextHandler")

try:
    from file_handlers.docx_handler import DocxHandler
except ImportError:
    print("警告：無法導入 DocxHandler")

try:
    from file_handlers.pdf_handler import PdfHandler
except ImportError:
    print("警告：無法導入 PdfHandler")

try:
    from docx import Document  # for .docx
except ImportError:
    Document = None

# Optional PDF support
try:
    from pypdf import PdfReader  # lightweight PDF text extraction
except Exception:
    PdfReader = None

# Optional PDF rasterizer (image previews)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# Optional .doc (Word 97-2003) via COM on Windows if MS Word is installed
try:
    import win32com.client  # type: ignore
except Exception:
    win32com = None  # noqa: N816

# 與前端相同的選項定義（測試版本僅用於輸出說明）
OPTION_META = {
    "name":       {"label": "姓名",     "desc": "姓名測試行"},
    "email":      {"label": "Email",    "desc": "Email 測試行"},
    "phone":      {"label": "電話",     "desc": "電話測試行"},
    "address":    {"label": "地址",     "desc": "地址測試行"},
    "birthday":   {"label": "生日",     "desc": "生日測試行"},
    "id":         {"label": "身分證",   "desc": "身分證測試行"},
    "student_id": {"label": "學號",     "desc": "學號測試行"},
    "org":        {"label": "機構",     "desc": "機構測試行"},
}

APP_ROOT = Path(__file__).resolve().parent
APP_OUTPUT_DIR = APP_ROOT / "test_output"
APP_PREVIEW_DIR = APP_OUTPUT_DIR / "_previews"
APP_OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
APP_EXPORT_DIR = APP_OUTPUT_DIR / "exports"  # ← 新增
APP_PREVIEW_DIR.mkdir(exist_ok=True, parents=True)

def _downloads_dir() -> Path:
    # 優先使用使用者的 Downloads；沒有就用系統暫存
    if os.name == "nt":
        home = Path(os.environ.get("USERPROFILE", str(Path.home())))
        dl = home / "Downloads"
    else:
        dl = Path.home() / "Downloads"
    return dl if dl.exists() else Path(tempfile.gettempdir())

def _text_to_pdf_no_font(src_txt_path: str, out_pdf_path: str) -> str:
    print("[PREVIEW] TXT->PDF via helv only:", src_txt_path)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    rect = fitz.Rect(36, 36, 559, 806)
    with open(src_txt_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    # 關鍵：絕不傳 fontfile，只用 fontname="helv"
    page.insert_textbox(rect, text, fontname="helv", fontsize=11, color=(0, 0, 0))
    doc.save(out_pdf_path)
    return out_pdf_path

def _docx_to_pdf_no_font(src_docx_path: str, out_pdf_path: str) -> str:
    print("[PREVIEW] DOCX->PDF try docx2pdf, fallback to helv:", src_docx_path)
    try:
        from docx2pdf import convert  # 可選
        convert(src_docx_path, out_pdf_path)
        return out_pdf_path
    except Exception:
        try:
            from docx import Document
            tmp_txt = APP_PREVIEW_DIR / (Path(src_docx_path).stem + "_plain.txt")
            with open(tmp_txt, "w", encoding="utf-8") as wf:
                doc = Document(src_docx_path)
                for p in doc.paragraphs:
                    wf.write(p.text + "\n")
                for t in doc.tables:
                    for row in t.rows:
                        wf.write(" | ".join(c.text for c in row.cells) + "\n")
            return _text_to_pdf_no_font(str(tmp_txt), out_pdf_path)
        except Exception as e:
            print("[PREVIEW] DOCX extract failed:", e)
            doc = fitz.open(); doc.new_page(); doc.save(out_pdf_path); return out_pdf_path

def _ensure_pdf_for_preview(processed_path: str) -> str:
    p = Path(processed_path)
    print("[PREVIEW] ensure PDF for:", processed_path)
    if p.suffix.lower() == ".pdf":
        return str(p)
    preview_pdf = APP_PREVIEW_DIR / f"{p.stem}_preview.pdf"
    preview_pdf.parent.mkdir(parents=True, exist_ok=True)
    if p.suffix.lower() == ".docx":
        return _docx_to_pdf_no_font(str(p), str(preview_pdf))
    else:
        return _text_to_pdf_no_font(str(p), str(preview_pdf))

def _rasterize_pdf(pdf_path: str, limit_pages: int = 10, dpi: int = 144):
    print("[PREVIEW] rasterize:", pdf_path)
    urls = []
    out_dir = APP_PREVIEW_DIR / (Path(pdf_path).stem + "_pages")
    shutil.rmtree(out_dir, ignore_errors=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    with fitz.open(pdf_path) as doc:
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for i, page in enumerate(doc):
            if i >= limit_pages:
                break
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = out_dir / f"page_{i+1:03d}.png"
            pix.save(img.as_posix())
            urls.append(img.as_uri())
    return urls


class TestBackend(QObject):
    """測試後端：接受 QML 上傳的檔案並回傳可顯示的預覽文字。

    介面完全比照 main.py 的 Backend，讓前端可以直接替換測試後端。
    支援 txt/md/log、docx、pdf，.doc 以 Windows COM（若可用）嘗試；否則回覆提示。
    """

    filesChanged = Signal(list)
    resultsReady = Signal(str)  # JSON: [{fileName, type, originalText, maskedText}]
    exportReady = Signal(str)   # ← 打包成功，回傳 zip 的 file:// URL
    exportFailed = Signal(str)  # ← 打包失敗訊息
    outputsCleared = Signal(str)       # 清理完成訊息
    outputsClearFailed = Signal(str)

    def __init__(self):
        super().__init__()
        self._files: list[str] = []
        self._options: list[str] = []
        self._last_results = []          # 新增：快取最近一次結果

    # 檔案操作 -------------------------------------------------
    @Slot(str)
    def addFile(self, path: str):
        if path and os.path.isfile(path) and path not in self._files:
            self._files.append(path)
            self.filesChanged.emit(self._files)

    @Slot('QStringList')
    def addFiles(self, paths):
        changed = False
        for p in paths:
            if p and os.path.isfile(p) and p not in self._files:
                self._files.append(p)
                changed = True
        if changed:
            self.filesChanged.emit(self._files)

    @Slot(str)
    def removeFile(self, path: str):
        if path in self._files:
            self._files.remove(path)
            self.filesChanged.emit(self._files)

    @Slot()
    def clearFiles(self):
        if self._files:
            self._files.clear()
            self.filesChanged.emit(self._files)

    @Slot(result='QStringList')
    def files(self):
        return self._files

    # 選項 -----------------------------------------------------
    @Slot('QStringList')
    def setOptions(self, opts):
        self._options = [o for o in opts if o in OPTION_META]

    @Slot(result='QStringList')
    def getOptions(self):
        return self._options

    @Slot(result=str)
    def getLastResults(self):
        """ResultPage 補抓"""
        return json.dumps(self._last_results, ensure_ascii=False)

    # 主流程 ---------------------------------------------------
    @Slot()
    def processFiles(self):
        """
        僅改『預覽』：先讓 Backend 產生去識別後檔案，再統一轉 PDF（不依賴外部字型），
        最後轉頁圖回傳給前端。核心處理流程不變。
        """
        from main import Backend
        backend = Backend()
        backend._options = self._options

        results = []
        for src in self._files:
            name = os.path.basename(src)
            if name.startswith("~$"):  # 跳過 Word 鎖定暫存檔
                continue
            ext = Path(src).suffix.lower()
            ftype = "pdf" if ext == ".pdf" else "docx" if ext == ".docx" else "text"

            try:
                out_path, _ = backend._process_file_with_deidentification(src, ftype, str(APP_OUTPUT_DIR))
                if not out_path or not os.path.isfile(out_path):
                    raise RuntimeError(f"後端未回傳有效輸出檔：{out_path}")

                pdf_path = _ensure_pdf_for_preview(out_path)   # 不用任何字型檔
                page_urls = _rasterize_pdf(pdf_path)
                results.append({
                    "fileName": os.path.basename(out_path),
                    "type": "pdf",
                    "originalText": "",
                    "maskedText": "",
                    "embedData": {
                        "viewType": "pdf",
                        "pageImages": page_urls,
                        "fileName": os.path.basename(out_path),
                        "pdfPath": Path(pdf_path).as_uri(),
                        "pageCount": len(page_urls),
                    },
                    "fileUrl": Path(out_path).as_uri(),
                })      
            except Exception as e:
                results.append({
                    "fileName": name, "type": ftype,
                    "originalText": "", "maskedText": "",
                    "embedData": {"viewType": "error", "error": str(e)},
                })
                        
        self._last_results = results               
        self.resultsReady.emit(json.dumps(results, ensure_ascii=False))

    # 打包全部處理後檔案成 ZIP
    @Slot()
    def exportAll(self):
        try:
            # 先用本次結果的 fileUrl 蒐集
            files = []
            for item in (self._last_results or []):
                url = item.get("fileUrl") if isinstance(item, dict) else None
                p = self._file_url_to_path(url) if url else None
                if p and os.path.isfile(p):
                    files.append(p)

            # 若本次結果為空，退回掃描 test_output（排除 _previews/exports）
            if not files:
                for p in APP_OUTPUT_DIR.rglob("*"):
                    if p.is_file() and "_previews" not in p.parts and "exports" not in p.parts:
                        files.append(str(p))

            # 去重
            files = list(dict.fromkeys(files))

            if not files:
                self.exportFailed.emit("沒有可匯出的檔案")
                return

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            zip_path = APP_EXPORT_DIR / f"results_{ts}.zip"

            manifest = []
            with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
                for fp in files:
                    arc = os.path.basename(fp)  # 壓縮包內僅保留檔名
                    z.write(fp, arcname=arc)
                    manifest.append({"file": arc, "src": fp})
                z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))

            self.exportReady.emit(Path(zip_path).as_uri())
        except Exception as e:
            self.exportFailed.emit(str(e))

    @Slot()
    def exportAllAndClear(self):
        """
        將 test_output 整個資料夾(遞迴)打包成 ZIP -> 放到 Downloads
        然後清空 test_output。成功會 emit exportReady(url)，失敗 emit exportFailed(msg)。
        """
        try:
            if not APP_OUTPUT_DIR.exists():
                self.exportFailed.emit("test_output 不存在")
                return

            # 建立 zip 目的地
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            dst_dir = _downloads_dir()
            zip_path = dst_dir / f"AnoniMe_results_{ts}.zip"

            # 打包（保留 test_output 內部相對路徑結構）
            with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
                base_len = len(str(APP_OUTPUT_DIR.parent)) + 1
                for root, dirs, files in os.walk(APP_OUTPUT_DIR):
                    for name in files:
                        fp = Path(root) / name
                        # 避免把未來可能放在 test_output 的 zip 也一起再打包
                        if fp.resolve() == zip_path.resolve():
                            continue
                        arcname = str(fp)[base_len:]  # 例如 test_output\processed\xxx
                        zf.write(fp, arcname)

            # 清空 test_output
            try:
                shutil.rmtree(APP_OUTPUT_DIR, ignore_errors=True)
            finally:
                APP_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

            # 回報成功（QML 收到後開啟並導回首頁）
            self.exportReady.emit(zip_path.as_uri())
            self.outputsCleared.emit("已清除 test_output")
        except Exception as e:
            self.exportFailed.emit(str(e))

    @Slot(bool)
    def clearTestOutput(self, keepExports: bool = False):
        """
        清除 test_output。keepExports 參數目前保留，預設 False=整個清掉。
        """
        try:
            if APP_OUTPUT_DIR.exists():
                shutil.rmtree(APP_OUTPUT_DIR, ignore_errors=True)
            APP_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            self.outputsCleared.emit("已清除 test_output")
        except Exception as e:
            self.outputsClearFailed.emit(str(e))

    # 讀取 -----------------------------------------------------
    def _read_file_preview(self, path: str, ftype: str, max_chars: int = 4000) -> str:
        try:
            if ftype == 'text':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    data = f.read(max_chars + 1024)  # 多讀少量以便截斷提示
                return self._truncate(data, max_chars) or "(空白)"

            if ftype == 'docx':
                if Document is None:
                    return "[缺少 python-docx 套件：pip install python-docx]"
                doc = Document(path)
                parts, length = [], 0
                for p in doc.paragraphs:
                    t = (p.text or '').strip()
                    if t:
                        parts.append(t)
                        length += len(t)
                    if length > max_chars:
                        break
                return self._truncate("\n".join(parts), max_chars) or "(空白 DOCX)"

            if ftype == 'doc':
                # 嘗試透過 Word COM（需要 Windows + 已安裝 Word + pywin32）
                if win32com is None:
                    return "[DOC 讀取需要 pywin32 以及安裝 Microsoft Word]"
                try:
                    word = win32com.gencache.EnsureDispatch('Word.Application')  # type: ignore[attr-defined]
                    word.Visible = False
                    doc = word.Documents.Open(path, ReadOnly=True)
                    text = doc.Content.Text if doc and doc.Content else ''
                    doc.Close(False)
                    word.Quit()
                    return self._truncate(text.replace("\r", "\n"), max_chars) or "(空白 DOC)"
                except Exception as e:
                    try:
                        # 盡力關閉殘留的 Word 實例
                        word.Quit()
                    except Exception:
                        pass
                    return f"[DOC 讀取失敗] {e}"

            if ftype == 'pdf':
                if PdfReader is None:
                    return "[缺少 pypdf 套件：pip install pypdf]"
                try:
                    reader = PdfReader(path)
                    parts, length = [], 0
                    for page in reader.pages:
                        try:
                            t = page.extract_text() or ''
                        except Exception:
                            t = ''
                        if t:
                            parts.append(t.strip())
                            length += len(t)
                        if length > max_chars:
                            break
                    return self._truncate("\n".join(parts), max_chars) or ""
                except Exception as e:
                    return f"[PDF 讀取錯誤] {e}"

            # 其他類型
            size = os.path.getsize(path)
            return f"[非文字類型，大小 {size} bytes]"
        except Exception as e:
            return f"[讀取錯誤] {e}"

    def _truncate(self, data: str, max_chars: int) -> str:
        if data is None:
            return ''
        if len(data) > max_chars:
            return data[:max_chars] + "\n...(截斷)"
        return data

    def _to_file_url(self, path: str) -> str:
        try:
            return Path(path).absolute().as_uri()
        except Exception:
            # 後備：手動組 file:///
            p = os.path.abspath(path).replace('\\', '/')
            if not p.startswith('/'):
                p = '/' + p
            return 'file://' + p

    def _render_pdf_pages(self, path: str, max_pages: int = 10, zoom: float = 1.5) -> list:
        """使用 PyMuPDF 產出多頁 PNG，回傳 file:// URL 陣列。無 PyMuPDF 則回空陣列。"""
        if fitz is None:
            return []
        try:
            doc = fitz.open(path)
            out_dir = tempfile.mkdtemp(prefix='AnoniMe_pdf_')
            urls = []
            page_count = min(len(doc), max_pages)
            mat = fitz.Matrix(zoom, zoom)
            for i in range(page_count):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_path = os.path.join(out_dir, f'page_{i+1}.png')
                pix.save(img_path)
                urls.append(Path(img_path).as_uri())
            doc.close()
            return urls
        except Exception:
            return []

    # 建立簡單測試版本（僅前置/後置說明） -----------------------
        """根據副檔名偵測語法類型，用於語法高亮提示。"""
        syntax_map = {
            '.py': 'python',
            '.js': 'javascript', 
            '.ts': 'typescript',
            '.html': 'html',
            '.css': 'css',
            '.json': 'json',
            '.xml': 'xml',
            '.yml': 'yaml',
            '.yaml': 'yaml',
            '.md': 'markdown',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bat': 'batch',
            '.ps1': 'powershell',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.java': 'java',
            '.cs': 'csharp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust'
        }
        return syntax_map.get(ext, 'text')

    def _create_hex_view(self, path: str, base_data: dict, max_bytes: int = 1024) -> dict:
        """為二進位檔案創建十六進制檢視。"""
        try:
            with open(path, 'rb') as f:
                data = f.read(max_bytes)
            
            hex_lines = []
            for i in range(0, len(data), 16):
                chunk = data[i:i+16]
                hex_part = ' '.join(f'{b:02X}' for b in chunk)
                ascii_part = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
                hex_lines.append({
                    "offset": f"{i:08X}",
                    "hex": hex_part,
                    "ascii": ascii_part
                })
            
            return {
                **base_data,
                "viewType": "hex",
                "hexLines": hex_lines,
                "totalBytes": len(data),
                "isPartial": len(data) >= max_bytes
            }
        except Exception:
            return {**base_data, "viewType": "error", "error": "無法讀取檔案"}

    # 新增：建立預設的處理後輸出路徑（系統暫存目錄/AnoniMe/processed）
    def _make_processed_output_path(self, in_path: str) -> str:
        base = os.path.splitext(os.path.basename(in_path))[0]
        ext = os.path.splitext(in_path)[1].lower()
        if ext not in ('.pdf', '.docx', '.txt'):
            # 預設轉回原副檔名
            out_ext = ext or '.out'
        else:
            out_ext = ext
        out_dir = os.path.join(tempfile.gettempdir(), 'AnoniMe', 'processed')
        os.makedirs(out_dir, exist_ok=True)
        return os.path.join(out_dir, f"{base}_deid{out_ext}")

    # 新增：同步執行處理腳本，擷取 stdout 中的 file:// URL 並回傳 (url, path)
    def _run_processor_and_get_output(self, script_name: str, in_path: str, out_path: str | None = None) -> tuple[str | None, str | None]:
        script_path = Path(__file__).resolve().parent / script_name
        if not script_path.exists():
            print(f"[processor missing] {script_name} 不存在")
            return None, None
        try:
            cmd = [sys.executable, str(script_path), in_path]
            if out_path:
                cmd.append(out_path)
            res = subprocess.run(cmd, cwd=str(script_path.parent), capture_output=True, text=True, timeout=600)
            stdout = (res.stdout or '').strip()
            stderr = (res.stderr or '').strip()
            if res.returncode != 0:
                print(f"[processor nonzero] {script_name} rc={res.returncode}")
                if stderr:
                    print(stderr[:2000])
            # 1) 優先：找最後一行 file:// URL
            url = None
            if stdout:
                lines = [ln.strip() for ln in stdout.splitlines() if ln.strip()]
                for ln in reversed(lines):
                    if ln.lower().startswith('file://'):
                        url = ln
                        break
                # 2) 次要：行中包含 file:// 片段
                if url is None:
                    m = re.search(r'file://\S+', stdout, flags=re.IGNORECASE)
                    if m:
                        url = m.group(0)
                # 3) 次要：嘗試從 stdout 解析 Windows 路徑
                if url is None:
                    # 取最後一個存在的檔案路徑
                    for ln in reversed(lines):
                        # 移除常見前綴（例如：輸出檔案: ）
                        cand = re.sub(r'^[^A-Za-z0-9]+', '', ln)
                        # 簡單偵測 Windows/Unix 路徑
                        win_path = re.search(r'[A-Za-z]:\\[^\n\r]+', cand)
                        if win_path and os.path.exists(win_path.group(0)):
                            p = win_path.group(0)
                            try:
                                return Path(p).resolve().as_uri(), p
                            except Exception:
                                return None, None
                        if os.path.exists(cand):
                            try:
                                return Path(cand).resolve().as_uri(), cand
                            except Exception:
                                return None, None
            # 4) 若指定 out_path 且檔案存在，直接回傳
            if url is None and out_path and os.path.exists(out_path):
                try:
                    return Path(out_path).resolve().as_uri(), out_path
                except Exception:
                    pass
            if not url:
                print('[processor parse] 未取得輸出 URL，stdout/stderr:')
                if stdout:
                    print(stdout[:2000])
                if stderr:
                    print(stderr[:2000])
                return None, None
            # 轉成本機路徑
            path = self._file_url_to_path(url)
            if path is None:
                # 如果不是 URL，視為本機路徑並回推 URL
                path = url
                try:
                    url = Path(path).resolve().as_uri()
                except Exception:
                    url = None
            return url, path
        except Exception as e:
            print(f"[processor exec error] {script_name}: {e}")
            return None, None

    # 新增：file:// URL 轉成本機路徑（含 Windows 相容）
    def _file_url_to_path(self, url: str) -> str | None:
        try:
            if not url.lower().startswith('file:'):
                return None
            u = urlparse(url)
            p = unquote(u.path)
            if sys.platform == 'win32':
                # 處理 /C:/ 形式
                if p.startswith('/') and len(p) > 3 and p[2] == ':':
                    p = p[1:]
                return p.replace('/', '\\')
            return p
        except Exception:
            return None

    def _convert_office_to_pdf(self, path: str) -> str | None:
        """使用 Word COM 將 DOC/DOCX 轉為 PDF。僅在 Windows + 安裝 Word + pywin32 時可行。"""
        if sys.platform != 'win32' or win32com is None:
            return None
        word = None
        try:
            word = win32com.client.Dispatch('Word.Application')  # type: ignore[attr-defined]
            word.Visible = False
            doc = word.Documents.Open(path)
            # 將 PDF 輸出到工作目錄的 tmp_preview
            out_dir = os.path.join(os.getcwd(), 'tmp_preview')
            os.makedirs(out_dir, exist_ok=True)
            out_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(path))[0] + '.pdf')
            wdFormatPDF = 17
            doc.SaveAs(out_pdf, FileFormat=wdFormatPDF)
            doc.Close(False)
            return out_pdf
        except Exception:
            return None
        finally:
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

    @Slot(result=str)
    def getTempOutputDir(self) -> str:
        """回傳應用程式臨時輸出資料夾（供 QML 顯示或開啟）。"""
        try:
            out_dir = os.path.join(tempfile.gettempdir(), 'AnoniMe')
            os.makedirs(out_dir, exist_ok=True)
            return out_dir
        except Exception as e:
            return f"[tempdir error] {e}"

    @Slot()
    def createTempTestFile(self):
        """建立一個臨時測試檔並把結果透過 resultsReady 發到 ResultPage。"""
        try:
            out_dir = self.getTempOutputDir()
            fn = f"preview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            p = os.path.join(out_dir, fn)
            with open(p, 'w', encoding='utf-8') as f:
                f.write(f"Temporary file created: {p}\n")
            payload = [{
                "fileName": os.path.basename(p),
                "type": "text",
                "originalText": f"Temporary file created: {p}",
                "maskedText": f"Temporary file created: {p}",
                "embedData": {"viewType": "text", "content": f"File path:\n{p}", "syntaxType": "text", "lineCount": 2},
                "fileUrl": Path(p).as_uri()
            }]
            self.resultsReady.emit(json.dumps(payload, ensure_ascii=False))
        except Exception as e:
            payload = [{
                "fileName": "temp_error.txt",
                "type": "text",
                "originalText": f"[error] {e}",
                "maskedText": f"[error] {e}",
                "embedData": {"viewType": "text", "content": str(e), "syntaxType": "text", "lineCount": 1}
            }]
            self.resultsReady.emit(json.dumps(payload, ensure_ascii=False))
