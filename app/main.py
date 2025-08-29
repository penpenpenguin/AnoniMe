import os, sys, json
from datetime import datetime
from PySide6.QtGui import QGuiApplication
from PySide6.QtQml import QQmlApplicationEngine
from PySide6.QtQuickControls2 import QQuickStyle
from PySide6.QtCore import QObject, Signal, Slot

# 嘗試引入檔案處理器，如果失敗則設為 None
try:
    from file_handlers.txt_handler import TextHandler
except ImportError:
    print("警告：無法導入 TextHandler")
    TextHandler = None

try:
    from file_handlers.docx_handler import DocxHandler
except ImportError:
    print("警告：無法導入 DocxHandler")
    DocxHandler = None

try:
    from file_handlers.pdf_handler import PdfHandler
except ImportError:
    print("警告：無法導入 PdfHandler")
    PdfHandler = None

try:
    from docx import Document
except ImportError:
    Document = None

# 追加：用於 PDF 轉圖與 Word 轉 PDF（可選）
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import win32com.client  # 需要安裝 pywin32 與本機有 Microsoft Word
except ImportError:
    win32com = None

import tempfile
import shutil
import subprocess  # 新增：用於呼叫 LibreOffice 無頭轉檔

# 勾選項目對應的顯示文字
OPTION_META = {
    "name":       {"label": "姓名",     "desc": "姓名測試行"},
    "email":      {"label": "Email",    "desc": "Email 測試行"},
    "phone":      {"label": "電話",     "desc": "電話測試行"},
    "address":    {"label": "地址",     "desc": "地址測試行"},
    "birthday":   {"label": "生日",     "desc": "生日測試行"},
    "id":         {"label": "身分證",   "desc": "身分證測試行"},
    "student_id": {"label": "學號",     "desc": "學號測試行"},
    "org":        {"label": "機構",     "desc": "機構測試行"},
}

class Backend(QObject):
    @Slot(str)
    def removeFile(self, path):
        """
        刪除指定路徑的檔案（供 UI 刪除功能使用）。
        """
        try:
            if path and os.path.isfile(path):
                os.remove(path)
                print(f"[刪除] 已移除檔案：{path}")
                # 從 _files 清單移除
                if path in self._files:
                    self._files.remove(path)
                return True
            else:
                print(f"[刪除失敗] 路徑無效或檔案不存在：{path}")
                return False
        except Exception as e:
            print(f"[刪除失敗] {e}")
            return False
    @Slot(str, str)
    def saveFileTo(self, src_path, dest_path):
        """
        將 src_path 檔案複製到 dest_path（用於下載功能）。
        """
        try:
            import shutil
            if src_path and dest_path and os.path.isfile(src_path):
                shutil.copy2(src_path, dest_path)
                print(f"[下載] 已將 {src_path} 複製到 {dest_path}")
                return True
            else:
                print(f"[下載失敗] 路徑無效或檔案不存在：{src_path} -> {dest_path}")
                return False
        except Exception as e:
            print(f"[下載失敗] {e}")
            return False
    @Slot(str, result=str)
    def readFileContent(self, path):
        """
        讀取指定檔案內容，支援 txt/docx，其他回傳占位文字。
        """
        # 支援 file:// URI -> 轉成本機路徑
        try:
            if isinstance(path, str) and path.startswith('file:'):
                from urllib.parse import urlparse, unquote
                p = urlparse(path)
                local = unquote(p.path)
                # Windows 路徑會以 /C:/ 開頭，去掉前導斜線
                if os.name == 'nt' and local.startswith('/') and len(local) > 2 and local[2] == ':':
                    local = local[1:]
                path = local
        except Exception:
            pass

        if not path or not os.path.isfile(path):
            return "[檔案不存在]"

        ext = path.lower().rsplit('.', 1)[-1] if '.' in path else ''
        try:
            if ext == "txt":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    data = f.read(8000)
                return data if data else "(空白)"
            if ext == "docx":
                if Document is None:
                    return "[缺少 python-docx 套件]"
                doc = Document(path)
                parts = []
                for p in doc.paragraphs:
                    parts.append(p.text)
                return "\n".join(parts) if parts else "(空白 DOCX)"
            if ext == "pdf":
                # PDF 不直接回傳完整文本；回傳簡短預覽與檔案路徑資訊
                return f"[PDF 檔案] {os.path.basename(path)}"
            return f"[不支援預覽: {ext}]"
        except Exception as e:
            return f"[讀取錯誤] {e}"
    filesChanged = Signal(list)
    resultsReady = Signal(str)  # JSON: [{fileName, type, originalText, maskedText}]

    def __init__(self):
        super().__init__()
        self._files = []
        self._options = []
        self._options_texts = []
        self._last_results = []          # 新增：快取最近一次結果

    # 檔案操作 -------------------------------------------------
    @Slot(str)
    def addFile(self, path: str):
        if path and os.path.isfile(path) and path not in self._files:
            self._files.append(path)
            self.filesChanged.emit(self._files)

    @Slot('QStringList')
    def addFiles(self, paths):
        changed = False
        for p in paths:
            if p and os.path.isfile(p) and p not in self._files:
                self._files.append(p)
                changed = True
        if changed:
            self.filesChanged.emit(self._files)

    @Slot(str)
    def removeFile(self, path: str):
        if path in self._files:
            self._files.remove(path)
            self.filesChanged.emit(self._files)

    @Slot()
    def clearFiles(self):
        if self._files:
            self._files.clear()
            self.filesChanged.emit(self._files)

    @Slot(result='QStringList')
    def files(self):
        return self._files

    # 選項 -----------------------------------------------------
    @Slot('QStringList')
    def setOptions(self, opts):
        self._options = [o for o in opts if o in OPTION_META]

    @Slot('QStringList')
    def setOptionsText(self, texts):
        """接收來自 UI 的選項顯示文字陣列（例如 ['姓名','Email']）。"""
        try:
            self._options_texts = list(texts) if texts is not None else []
            print("Backend: setOptionsText ->", self._options_texts)
        except Exception as e:
            print("setOptionsText error:", e)

    @Slot(result='QStringList')
    def getOptions(self):
        return self._options

    # 主流程 ---------------------------------------------------
    @Slot()
    def processFiles(self):
        print("Backend: processFiles() start, files =", len(self._files), "options =", self._options)
        results = []
        
        # 創建輸出目錄
        output_dir = os.path.join(os.getcwd(), "test_output")
        os.makedirs(output_dir, exist_ok=True)
        
        for path in self._files:
            file_name = os.path.basename(path)
            ext = file_name.lower().rsplit('.', 1)[-1] if '.' in file_name else ''
            if ext in ('txt','md','log'):
                ftype = 'text'
            elif ext == 'docx':
                ftype = 'docx'
            elif ext == 'doc':
                ftype = 'doc'
            elif ext == 'pdf':
                ftype = 'pdf'
            else:
                ftype = 'binary'

            # 讀取原始檔案內容作為預覽
            original = self._read_file_preview(path, ftype)
            
            # 執行實際的去識別化處理
            masked_file_path, masked_text = self._process_file_with_deidentification(path, ftype, output_dir)

            # 產生內嵌檢視資料（使用處理後的檔案，如果存在的話）
            display_path = masked_file_path if masked_file_path else path
            embed_data = self._create_embed_data(display_path, ftype, file_name)

            results.append({
                "fileName": file_name,
                "type": ftype,
                "originalText": original,
                "maskedText": masked_text,
                "embedData": embed_data,
                "outputPath": masked_file_path
            })

        self._last_results = results[:]   # 快取
        print("Backend: emit resultsReady count =", len(results))
        self.resultsReady.emit(json.dumps(results, ensure_ascii=False))

    @Slot(result=str)
    def getLastResults(self):
        """ResultPage 補抓"""
        return json.dumps(self._last_results, ensure_ascii=False)

    # 讀取 -----------------------------------------------------
    def _read_file_preview(self, path, ftype, max_chars=4000):
        try:
            if ftype == 'text':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    data = f.read(max_chars + 1)
                if len(data) > max_chars:
                    data = data[:max_chars] + "\n...(截斷)"
                return data or "(空白)"
            if ftype == 'docx':
                if Document is None:
                    return "[缺少 python-docx 套件：pip install python-docx]"
                doc = Document(path)
                parts, length = [], 0
                
                for p in doc.paragraphs:
                    # 保持段落格式
                    paragraph_text = ""
                    
                    # 處理段落中的每個 run（文字片段）
                    for run in p.runs:
                        text = run.text or ''
                        if text:
                            # 根據格式添加標記
                            if run.bold:
                                text = f"**{text}**"
                            if run.italic:
                                text = f"*{text}*"
                            if run.underline:
                                text = f"_{text}_"
                        paragraph_text += text
                    
                    # 處理段落樣式
                    if paragraph_text.strip():
                        style_name = p.style.name if p.style else "Normal"
                        
                        # 根據樣式添加格式標記
                        if "Heading" in style_name:
                            level = 1
                            if "Heading 1" in style_name:
                                level = 1
                            elif "Heading 2" in style_name:
                                level = 2
                            elif "Heading 3" in style_name:
                                level = 3
                            elif "Heading 4" in style_name:
                                level = 4
                            elif "Heading 5" in style_name:
                                level = 5
                            elif "Heading 6" in style_name:
                                level = 6
                            paragraph_text = "#" * level + " " + paragraph_text.strip()
                        elif "List" in style_name or "Bullet" in style_name:
                            paragraph_text = "• " + paragraph_text.strip()
                        elif "Quote" in style_name:
                            paragraph_text = "> " + paragraph_text.strip()
                        
                        parts.append(paragraph_text)
                        length += len(paragraph_text)
                        
                        if length > max_chars:
                            break
                
                # 處理表格
                for table in doc.tables:
                    if length > max_chars:
                        break
                    
                    table_text = "\n[表格]\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_text += row_text + "\n"
                    
                    parts.append(table_text)
                    length += len(table_text)
                
                result = "\n\n".join(parts) if parts else "(空白 DOCX)"
                if len(result) > max_chars:
                    result = result[:max_chars] + "\n...(截斷)"
                return result
            if ftype == 'pdf':
                return f"[PDF 預覽占位] {os.path.basename(path)}"
            size = os.path.getsize(path)
            return f"[非文字類型，大小 {size} bytes]"
        except Exception as e:
            return f"[讀取錯誤] {e}"

    # 建立簡單測試版本（僅前置/後置說明） -----------------------
    def _build_test_version(self, original: str, filename: str) -> str:
        if not original:
            return original
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        header = [
            "[測試版本輸出]",
            f"[檔案: {filename}]",
            f"[時間: {now}]",
            f"[原文字元數: {len(original)}]"
        ]
        if self._options:
            header.append("[已勾選項目] " + ", ".join(OPTION_META[o]['label'] for o in self._options))
            header.append("--- 選項測試行開始 ---")
            for o in self._options:
                meta = OPTION_META[o]
                # 直接一行，不做任何內文搜尋 / 替換
                header.append(f"[TEST-{o.upper()}] {meta['desc']} (未修改正文)")
            header.append("--- 選項測試行結束 ---")
        else:
            header.append("[未勾選任何項目]")

        footer = [
            "",
            "------ TEST SUMMARY ------",
            f"選項數: {len(self._options)}",
            "本測試版本僅在前方加入描述，正文保持原樣。",
            "--------------------------"
        ]
        return "\n".join(header) + "\n\n" + original + "\n\n" + "\n".join(footer)

    # 實際去識別化處理函數 ------------------------------------
    def _process_file_with_deidentification(self, input_path: str, ftype: str, output_dir: str):
        """
        根據檔案類型將檔案路徑路由到 scripts 目錄中的對應處理器
        回傳: (處理後檔案路徑, 處理後文字內容的預覽)
        """
        from pathlib import Path
        os.makedirs(output_dir, exist_ok=True)
        base = Path(input_path).stem

        if ftype == "text":
            # return self._process_txt_file(input_path, output_dir)
            out_path = os.path.join(output_dir, f"{base}_deid.txt")
            handler = TextHandler()
            handler.deidentify(input_path, out_path, language="auto")
            return out_path, ""

        if ftype == "docx":
            out_path = os.path.join(output_dir, f"{base}_deid.docx")
            handler = DocxHandler()
            handler.deidentify(input_path, out_path, language="auto")
            return out_path, ""

        if ftype == "pdf":
            out_path = os.path.join(output_dir, f"{base}_deid.pdf")
            detector = PdfHandler()
            detector.deidentify(input_path, out_path, language="auto")
            return out_path, ""

        # 其他副檔名...
        return None, f"[不支援的檔案類型: {ftype}]"

    def _process_txt_file(self, input_path: str, output_path: str):
        """處理文字檔案"""
        try:
            if TextHandler is None:
                return self._fallback_copy_file(input_path, output_path, "TextHandler 不可用")
            
            handler = TextHandler()
            result_path = handler.deidentify(input_path=input_path, output_path=output_path)
            
            # 讀取處理後的內容作為預覽
            with open(result_path, 'r', encoding='utf-8', errors='ignore') as f:
                masked_content = f.read(4000)  # 限制預覽長度
            
            print(f"✅ TXT 去識別化完成！輸出檔案：{result_path}")
            return result_path, masked_content
            
        except Exception as e:
            return self._fallback_copy_file(input_path, output_path, f"TXT 處理失敗: {str(e)}")
    
    def _process_docx_file(self, input_path: str, output_path: str):
        """處理 Word 檔案"""
        try:
            if DocxHandler is None:
                return self._fallback_copy_file(input_path, output_path, "DocxHandler 不可用")
            
            handler = DocxHandler()
            result_path = handler.deidentify(input_path=input_path, output_path=output_path)
            
            # 讀取處理後的內容作為預覽
            masked_content = self._read_file_preview(result_path, 'docx')
            
            print(f"✅ DOCX 去識別化完成！輸出檔案：{result_path}")
            return result_path, masked_content
            
        except Exception as e:
            return self._fallback_copy_file(input_path, output_path, f"DOCX 處理失敗: {str(e)}")
    
    def _process_pdf_file(self, input_path: str, output_path: str):
        """處理 PDF 檔案"""
        try:
            if PdfHandler is None:
                return self._fallback_copy_file(input_path, output_path, "PdfHandler 不可用")
            
            detector = PdfHandler()
            result_path = detector.deidentify(input_path=input_path, output_path=output_path)
            
            # PDF 預覽
            masked_content = f"[PDF 去識別化完成] {os.path.basename(result_path)}"
            
            print(f"✅ PDF 去識別化完成！輸出檔案：{result_path}")
            return result_path, masked_content
            
        except Exception as e:
            return self._fallback_copy_file(input_path, output_path, f"PDF 處理失敗: {str(e)}")
    
    def _fallback_copy_file(self, input_path: str, output_path: str, reason: str):
        """
        當處理器不可用時的回退方案：直接複製檔案
        """
        try:
            import shutil
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(input_path, output_path)
            original_content = self._read_file_preview(input_path, os.path.splitext(input_path)[1][1:])
            print(f"⚠️ {reason}，已複製原始檔案到：{output_path}")
            return output_path, f"[{reason}，已複製原始檔案]\n{original_content}"
        except Exception as e:
            original_content = self._read_file_preview(input_path, os.path.splitext(input_path)[1][1:])
            return None, f"[複製失敗: {str(e)}]\n{original_content}"

    def _create_embed_data(self, path: str, ftype: str, file_name: str):
        try:
            if ftype in ('doc', 'docx'):
                # 嘗試以 Word 轉 PDF（保持版面），若不可用再以 LibreOffice 無頭轉 PDF
                pdf_path = self._convert_office_to_pdf(path)
                if pdf_path:
                    pages, meta = self._render_pdf_pages(pdf_path)
                    if pages:
                        return {
                            "viewType": "pdf",
                            "pageImages": pages,
                            "pageCount": len(pages),
                            "metadata": meta,
                            "fileName": file_name
                        }
                    return {"viewType": "error", "error": "PDF 轉圖失敗", "fileName": file_name}
                return {
                    "viewType": "unsupported",
                    "reason": "無法將 Word 轉為 PDF；請安裝 Microsoft Word（或安裝 LibreOffice 並設定 PATH）",
                    "fileName": file_name
                }

            if ftype == 'pdf':
                pages, meta = self._render_pdf_pages(path)
                if pages:
                    return {
                        "viewType": "pdf",
                        "pageImages": pages,
                        "pageCount": len(pages),
                        "metadata": meta,
                        "fileName": file_name
                    }
                return {"viewType": "error", "error": "PDF 轉圖失敗", "fileName": file_name}

            if ftype == 'text':
                content = self._read_file_preview(path, 'text', 8000)
                return {
                    "viewType": "text",
                    "content": content,
                    "syntaxType": "text",
                    "lineCount": self._estimate_line_count_of_file(path)
                }
        except Exception as e:
            return {"viewType": "error", "error": str(e), "fileName": file_name}
        return {}

    def _render_pdf_pages(self, pdf_path: str):
        if fitz is None:
            return [], {}
        try:
            doc = fitz.open(pdf_path)
            # 圖片輸出資料夾（工作目錄下 tmp_preview）
            out_dir = os.path.join(os.getcwd(), 'tmp_preview')
            os.makedirs(out_dir, exist_ok=True)
            page_paths = []
            meta = {}
            try:
                meta_raw = doc.metadata or {}
                meta = {"title": meta_raw.get("title"), "author": meta_raw.get("author")}
            except Exception:
                meta = {}

            zoom = 2.0  # 提升解析度，減少鋸齒與條紋
            mat = fitz.Matrix(zoom, zoom)
            base = os.path.splitext(os.path.basename(pdf_path))[0]
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_path = os.path.join(out_dir, f"{base}_page_{i+1}.png")
                pix.save(img_path)
                page_paths.append(img_path)
            doc.close()
            return page_paths, meta
        except Exception:
            return [], {}

    def _convert_office_to_pdf(self, path: str):
        """優先使用 Word COM；失敗時嘗試 LibreOffice soffice 無頭轉 PDF。"""
        # 先試 Word COM（Windows + Word + pywin32）
        if sys.platform == 'win32' and win32com is not None:
            word = None
            try:
                word = win32com.client.Dispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(path)
                tmp_dir = os.path.join(os.getcwd(), 'tmp_preview')
                os.makedirs(tmp_dir, exist_ok=True)
                out_pdf = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + '.pdf')
                wdFormatPDF = 17
                doc.SaveAs(out_pdf, FileFormat=wdFormatPDF)
                doc.Close(False)
                return out_pdf if os.path.exists(out_pdf) else None
            except Exception:
                # 失敗則繼續試 LibreOffice
                pass
            finally:
                try:
                    if word is not None:
                        word.Quit()
                except Exception:
                    pass
        
        # 再試 LibreOffice soffice
        return self._convert_office_to_pdf_with_soffice(path)

    def _convert_office_to_pdf_with_soffice(self, path: str):
        """使用 LibreOffice 無頭模式將 DOC/DOCX 轉為 PDF。成功返回 PDF 路徑，否則返回 None。"""
        soffice = self._find_soffice_path()
        if not soffice:
            return None
        try:
            out_dir = os.path.join(os.getcwd(), 'tmp_preview')
            os.makedirs(out_dir, exist_ok=True)
            # soffice --headless --convert-to pdf --outdir <out_dir> <path>
            cmd = [soffice, "--headless", "--nologo", "--convert-to", "pdf", "--outdir", out_dir, path]
            proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
            if proc.returncode != 0:
                return None
            out_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(path))[0] + '.pdf')
            return out_pdf if os.path.exists(out_pdf) else None
        except Exception:
            return None

    def _find_soffice_path(self) -> str | None:
        """嘗試尋找 LibreOffice soffice 可執行檔。可透過環境變數 SOFFICE_PATH 指定。"""
        # 1) 環境變數
        p = os.environ.get('SOFFICE_PATH')
        if p and os.path.exists(p):
            return p
        # 2) 嘗試系統 PATH
        try:
            return shutil.which('soffice')
        except Exception:
            pass
        # 3) 常見安裝路徑（Windows）
        candidates = [
            r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe"
        ]
        for c in candidates:
            if os.path.exists(c):
                return c
        return None

    def _estimate_line_count_of_file(self, path: str, max_lines: int = 2000):
        try:
            cnt = 0
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                for _ in f:
                    cnt += 1
                    if cnt >= max_lines:
                        break
            return cnt
        except Exception:
            return 0

# 啟動 --------------------------------------------------------
if __name__ == "__main__":
    QQuickStyle.setStyle("Fusion")
    app = QGuiApplication(sys.argv)
    engine = QQmlApplicationEngine()
    backend = Backend()
    engine.rootContext().setContextProperty("backend", backend)
    # Ensure we load the Main.qml from the project's qml directory
    qml_file = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'qml', 'Main.qml')
    qml_url = 'file:///' + qml_file.replace('\\', '/')
    print(f"Loading QML from: {qml_url}")
    engine.load(qml_url)
    if not engine.rootObjects():
        print("QQmlApplicationEngine failed to load. Check QML error output above.")
        sys.exit(-1)
    sys.exit(app.exec())