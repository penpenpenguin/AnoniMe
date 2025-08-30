import os
import json
import shutil
from datetime import datetime
import tempfile
from pathlib import Path
import sys  # 新增：供平台判斷
import subprocess  # 新增：啟動外部處理腳本
from urllib.parse import urlparse, unquote  # 新增：解析 file:// URL
import re  # 新增：解析 stdout 中的路徑
from zipfile import ZipFile, ZIP_DEFLATED

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from PySide6.QtCore import QObject, Signal, Slot
import traceback

try:
    from file_handlers.txt_handler import TextHandler
except ImportError:
    print("警告：無法導入 TextHandler")
    traceback.print_exc()

try:
    from file_handlers.docx_handler import DocxHandler
except ImportError:
    print("警告：無法導入 DocxHandler")
    traceback.print_exc()

try:
    from file_handlers.pdf_handler import PdfHandler
except ImportError:
    print("警告：無法導入 PdfHandler")
    traceback.print_exc()

# Optional packages for file preview
try:
    from docx import Document  # for .docx
except ImportError:
    Document = None

# Optional PDF support
try:
    from pypdf import PdfReader  # lightweight PDF text extraction
except Exception:
    PdfReader = None

# Optional PDF rasterizer (image previews)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# Optional .doc (Word 97-2003) via COM on Windows if MS Word is installed
try:
    import win32com.client  # type: ignore
except Exception:
    win32com = None  # noqa: N816

def _downloads_dir() -> Path:
    # 優先使用使用者的 Downloads；沒有就用系統暫存
    if os.name == "nt":
        home = Path(os.environ.get("USERPROFILE", str(Path.home())))
        dl = home / "Downloads"
    else:
        dl = Path.home() / "Downloads"
    return dl if dl.exists() else Path(tempfile.gettempdir())

APP_ROOT = Path(__file__).resolve().parent
APP_OUTPUT_DIR = APP_ROOT / "test_output"
APP_PREVIEW_DIR = APP_OUTPUT_DIR / "_previews"
APP_OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
APP_EXPORT_DIR = APP_OUTPUT_DIR / "exports"  # ← 新增
APP_PREVIEW_DIR.mkdir(exist_ok=True, parents=True)

def _rasterize_pdf(pdf_path: str, limit_pages: int = 10, dpi: int = 144):
    print("[PREVIEW] rasterize:", pdf_path)
    urls = []
    out_dir = APP_PREVIEW_DIR / (Path(pdf_path).stem + "_pages")
    shutil.rmtree(out_dir, ignore_errors=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    with fitz.open(pdf_path) as doc:
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for i, page in enumerate(doc):
            if i >= limit_pages:
                break
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = out_dir / f"page_{i+1:03d}.png"
            pix.save(img.as_posix())
            urls.append(img.as_uri())
    return urls


class TestBackend(QObject):
    """正式後端：接受 QML 上傳的檔案與選項，回傳真實處理結果。

    介面完全比照 main.py 的 Backend，直接處理前端傳遞的檔案與選項，回傳 originalText、maskedText 等真實資料。
    """

    filesChanged = Signal(list)
    resultsReady = Signal(str)  # JSON: [{fileName, type, originalText, maskedText}]
    exportReady = Signal(str)   # ← 打包成功，回傳 zip 的 file:// URL
    exportFailed = Signal(str)  # ← 打包失敗訊息
    outputsCleared = Signal(str)       # 清理完成訊息
    outputsClearFailed = Signal(str)
    stateCleared = Signal()        

    def __init__(self):
        super().__init__()
        self._files: list[str] = []
        self._options: list[str] = []
        self._option_texts: list[str] = []    # 新增：儲存選項的顯示文字
        self._last_results = []          # 新增：快取最近一次結果

    # 檔案操作 -------------------------------------------------
    @Slot(str)
    def addFile(self, path: str):
        if path and os.path.isfile(path) and path not in self._files:
            self._files.append(path)
            self.filesChanged.emit(self._files)

    @Slot('QStringList')
    def addFiles(self, paths):
        changed = False
        for p in paths:
            if p and os.path.isfile(p) and p not in self._files:
                self._files.append(p)
                changed = True
        if changed:
            self.filesChanged.emit(self._files)

    @Slot(str)
    def removeFile(self, path: str):
        if path in self._files:
            self._files.remove(path)
            self.filesChanged.emit(self._files)

    @Slot()
    def clearFiles(self):
        if self._files:
            self._files.clear()
            self.filesChanged.emit(self._files)

    @Slot(result='QStringList')
    def files(self):
        return self._files

    # 選項 -----------------------------------------------------
    @Slot('QStringList')
    def setOptions(self, opts):
        print(f"[後端] setOptions 接收到: {opts}")  # 參考點，確認接收 list
        self._options = opts

    @Slot(result='QStringList')
    def getOptions(self):
        return self._options

    @Slot(str, result=str)
    def getFilePreviewData(self, filePath):
        """取得檔案的內嵌預覽資料，回傳 JSON 字串"""
        try:
            import json
            
            # 如果是 file:// URL，先轉換為本地路徑
            if filePath.startswith('file://'):
                actualPath = self._file_url_to_path(filePath)
                if not actualPath:
                    return json.dumps({"viewType": "error", "error": "無效的檔案 URL"})
            else:
                actualPath = filePath
            
            if not actualPath or not os.path.exists(actualPath):
                return json.dumps({"viewType": "error", "error": f"檔案不存在: {actualPath}"})
            
            # 取得檔案副檔名
            ext = os.path.splitext(actualPath)[1].lower().lstrip('.')
            if not ext:
                return json.dumps({"viewType": "error", "error": "無法判斷檔案類型"})
            
            # 根據副檔名判斷檔案類型
            if ext == "pdf":
                ftype = "pdf"
            elif ext == "docx":
                ftype = "docx"
            elif ext in ["txt", "csv", "html", "json", "md"]:
                ftype = "text"
            else:
                # 預設當作文字檔處理
                ftype = "text"
            
            print(f"[DEBUG] Preview for {actualPath}, ext={ext}, ftype={ftype}")
            
            # 使用現有的 _generate_embed_data 方法
            preview_data = self._generate_embed_data(actualPath, ftype)
            return json.dumps(preview_data)
            
        except Exception as e:
            import json
            import traceback
            traceback.print_exc()
            return json.dumps({"viewType": "error", "error": f"預覽生成失敗: {str(e)}"})

    def _generate_embed_data(self, file_path: str, ftype: str) -> dict:
        """根據檔案類型生成對應的 embedData，統一轉換為 PDF 預覽格式"""
        try:
            # 統一轉換為 PDF 預覽格式
            if ftype == "pdf":
                # PDF 檔案：直接生成頁面圖像
                if fitz is None:
                    return {"viewType": "error", "error": "PyMuPDF 未安裝，無法預覽 PDF"}
                
                try:
                    page_images = self._render_pdf_pages(file_path)
                    return {
                        "viewType": "pdf",
                        "pageImages": page_images,
                        "pageCount": len(page_images)
                    }
                except Exception as e:
                    return {"viewType": "error", "error": f"PDF 預覽生成失敗: {e}"}
                    
            elif ftype == "docx" or ftype == "text":
                # DOCX 和 TXT 檔案：先轉換為 PDF，再生成預覽
                if fitz is None:
                    return {"viewType": "error", "error": "PyMuPDF 未安裝，無法預覽檔案"}
                
                try:
                    # 嘗試轉換為 PDF
                    pdf_path = self._convert_to_pdf_for_preview(file_path, ftype)
                    if pdf_path and os.path.exists(pdf_path):
                        page_images = self._render_pdf_pages(pdf_path)
                        return {
                            "viewType": "pdf",
                            "pageImages": page_images,
                            "pageCount": len(page_images),
                            "originalType": ftype
                        }
                    else:
                        # 如果轉換失敗，回退到文字預覽
                        return self._generate_text_fallback(file_path, ftype)
                except Exception as e:
                    return {"viewType": "error", "error": f"檔案轉換預覽失敗: {e}"}
            
            else:
                return {"viewType": "error", "error": f"不支援的檔案類型: {ftype}"}
                
        except Exception as e:
            return {"viewType": "error", "error": f"生成預覽資料失敗: {e}"}

    def _convert_to_pdf_for_preview(self, file_path: str, ftype: str) -> str:
        """將檔案轉換為 PDF 用於預覽"""
        try:
            preview_dir = APP_PREVIEW_DIR / "converted_pdfs"
            preview_dir.mkdir(exist_ok=True, parents=True)
            
            base_name = Path(file_path).stem
            pdf_path = preview_dir / f"{base_name}_preview.pdf"
            
            if ftype == "docx":
                # 使用 Word COM 轉換 DOCX 為 PDF
                pdf_result = self._convert_office_to_pdf(file_path)
                if pdf_result and os.path.exists(pdf_result):
                    # 移動到預覽目錄
                    shutil.move(pdf_result, str(pdf_path))
                    return str(pdf_path)
                
            elif ftype == "text":
                # 將文字檔案轉換為簡單的 PDF
                return self._convert_text_to_pdf(file_path, str(pdf_path))
                
            return None
            
        except Exception as e:
            print(f"轉換為 PDF 失敗: {e}")
            return None
    
    def _convert_text_to_pdf(self, text_path: str, pdf_path: str) -> str:
        """將文字檔案轉換為簡單的 PDF"""
        try:
            # 這裡可以使用 reportlab 或其他庫來轉換
            # 暫時返回 None，表示不支援
            return None
        except Exception:
            return None
    
    def _generate_text_fallback(self, file_path: str, ftype: str) -> dict:
        """當 PDF 轉換失敗時的文字回退預覽"""
        try:
            if ftype == "text":
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                return {
                    "viewType": "text",
                    "content": content,
                    "syntaxType": self._detect_syntax_type(Path(file_path).suffix.lower()),
                    "lineCount": content.count('\n') + 1 if content else 1
                }
                
            elif ftype == "docx":
                if Document is None:
                    return {"viewType": "error", "error": "python-docx 未安裝"}
                
                doc = Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                content = "\n".join(text_parts)
                return {
                    "viewType": "text",
                    "content": content,
                    "syntaxType": "text",
                    "lineCount": len(text_parts)
                }
                
            return {"viewType": "error", "error": "無法生成預覽"}
            
        except Exception as e:
            return {"viewType": "error", "error": f"文字回退預覽失敗: {e}"}

    def _detect_syntax_type(self, ext: str) -> str:
        """根據副檔名偵測語法類型，用於語法高亮提示。"""
        syntax_map = {
            '.py': 'python',
            '.js': 'javascript', 
            '.ts': 'typescript',
            '.html': 'html',
            '.css': 'css',
            '.json': 'json',
            '.xml': 'xml',
            '.yml': 'yaml',
            '.yaml': 'yaml',
            '.md': 'markdown',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bat': 'batch',
            '.ps1': 'powershell',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.java': 'java',
            '.cs': 'csharp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust'
        }
        return syntax_map.get(ext, 'text')

    def _render_pdf_pages(self, path: str, max_pages: int = 10, zoom: float = 1.5) -> list:
        """使用 PyMuPDF 產出多頁 PNG，回傳 file:// URL 陣列。無 PyMuPDF 則回空陣列。"""
        if fitz is None:
            return []
        try:
            doc = fitz.open(path)
            out_dir = tempfile.mkdtemp(prefix='AnoniMe_pdf_')
            urls = []
            page_count = min(len(doc), max_pages)
            mat = fitz.Matrix(zoom, zoom)
            for i in range(page_count):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_path = os.path.join(out_dir, f'page_{i+1}.png')
                pix.save(img_path)
                urls.append(Path(img_path).as_uri())
            doc.close()
            return urls
        except Exception:
            return []

    # @Slot('QStringList')
    # def setOptionsText(self, optionTexts):
    #     """接收前端傳來的選項顯示文字列表"""
    #     print(f"[後端] setOptionsText 接收到: {optionTexts}")
    #     self._option_texts = list(optionTexts)
    #     print(f"[後端] 已儲存選項文字: {self._option_texts}")

    # @Slot(result='QStringList')
    # def getOptionsText(self):
    #     """回傳目前儲存的選項顯示文字列表"""
    #     return self._option_texts

    @Slot(result=str)
    def getLastResults(self):
        """ResultPage 補抓"""
        return json.dumps(self._last_results, ensure_ascii=False)

    # 主流程 ---------------------------------------------------
    @Slot()
    def processFiles(self):
        """
        正式後端：依據前端傳遞的檔案與選項，回傳真實處理結果。
        根據檔案類型分配給相對應的 handler 進行處理。
        """
        results = []
        for src in self._files:
            name = os.path.basename(src)
            if name.startswith("~$"):
                continue
            ext = Path(src).suffix.lower()
            
            # 根據副檔名判斷檔案類型
            if ext == ".pdf":
                ftype = "pdf"
            elif ext == ".docx":
                ftype = "docx"
            elif ext in [".txt", ".csv", ".html", ".json", ".md"]:
                ftype = "text"
            else:
                # 預設當作文字檔處理
                ftype = "text"
                print(f"[後端] 未知副檔名 {ext}，當作文字檔處理")

            try:
                # 建立輸出路徑到 test_output/processed
                processed_dir = APP_OUTPUT_DIR / "processed"
                processed_dir.mkdir(parents=True, exist_ok=True)
                out_filename = f"{Path(src).stem}_deid{ext if ext else '.txt'}"
                out_path = processed_dir / out_filename

                # 根據檔案類型選擇對應的 handler 並進行處理
                handler = None
                
                if ftype == "text":
                    try:
                        handler = TextHandler()
                        print(f"[後端] 使用 TextHandler 處理文字檔案: {src}")
                    except NameError:
                        raise RuntimeError("TextHandler 未正確導入或不可用")
                        
                elif ftype == "docx":
                    try:
                        handler = DocxHandler()
                        print(f"[後端] 使用 DocxHandler 處理 Word 檔案: {src}")
                    except NameError:
                        raise RuntimeError("DocxHandler 未正確導入或不可用")
                        
                elif ftype == "pdf":
                    try:
                        handler = PdfHandler()
                        print(f"[後端] 使用 PdfHandler 處理 PDF 檔案: {src}")
                    except NameError:
                        raise RuntimeError("PdfHandler 未正確導入或不可用")
                        
                else:
                    raise RuntimeError(f"不支援的檔案類型：{ftype} (副檔名: {ext})")

                # 確認 handler 已正確初始化
                if handler is None:
                    raise RuntimeError(f"無法取得 {ftype} 類型的處理器")

                # 呼叫 handler 的 deidentify 方法進行處理
                print(f"[後端] 開始去識別化處理，輸入: {src}, 輸出: {out_path}")
                # users options
                backend = TestBackend()
                selected_types_list = backend.getOptions()
                processed_path = handler.deidentify(src, str(out_path), selected_types_list)

                if not processed_path or not os.path.isfile(processed_path):
                    raise RuntimeError(f"Handler 未產生有效輸出檔：{processed_path}")

                print(f"[後端] 去識別化完成，輸出檔案: {processed_path}")

                results.append({
                    "fileName": os.path.basename(processed_path),
                    "type": ftype,
                    "outputPath": processed_path,  # ← 這裡要是 txt/docx/pdf 檔案路徑
                    "fileUrl": Path(processed_path).as_uri(),  # ← 這裡也是處理後檔案，不是 json

                })
                
                print(f"[後端] 檔案處理成功: {name} -> {os.path.basename(processed_path)}")

            except Exception as e:
                print(f"[後端] 檔案處理失敗: {name}, 錯誤: {e}")
                results.append({
                    "fileName": name,
                    "type": ftype,
                    "originalText": "",
                    "maskedText": "",
                    "error": str(e),
                })

        self._last_results = results
        print(f"[後端] 處理完成，共 {len(results)} 個檔案")
        self.resultsReady.emit(json.dumps(results, ensure_ascii=False))

    # 打包全部處理後檔案成 ZIP
    @Slot()
    def exportAll(self):
        try:
            # 先用本次結果的 fileUrl 蒐集
            files = []
            for item in (self._last_results or []):
                url = item.get("fileUrl") if isinstance(item, dict) else None
                p = self._file_url_to_path(url) if url else None
                if p and os.path.isfile(p):
                    files.append(p)

            # 若本次結果為空，退回掃描 test_output（排除 _previews/exports）
            if not files:
                for p in APP_OUTPUT_DIR.rglob("*"):
                    if p.is_file() and "_previews" not in p.parts and "exports" not in p.parts:
                        files.append(str(p))

            # 去重
            files = list(dict.fromkeys(files))

            if not files:
                self.exportFailed.emit("沒有可匯出的檔案")
                return

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            zip_path = APP_EXPORT_DIR / f"results_{ts}.zip"

            manifest = []
            with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
                for fp in files:
                    arc = os.path.basename(fp)  # 壓縮包內僅保留檔名
                    z.write(fp, arcname=arc)
                    manifest.append({"file": arc, "src": fp})
                z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))

            self.exportReady.emit(Path(zip_path).as_uri())
        except Exception as e:
            self.exportFailed.emit(str(e))

    @Slot(bool)
    def clearTestOutput(self, keepExports: bool = False):
        """
        清除 test_output。keepExports 參數目前保留，預設 False=整個清掉。
        """
        try:
            if APP_OUTPUT_DIR.exists():
                shutil.rmtree(APP_OUTPUT_DIR, ignore_errors=True)
            APP_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            self.outputsCleared.emit("已清除 test_output")
        except Exception as e:
            self.outputsClearFailed.emit(str(e))

    # 讀取 -----------------------------------------------------
    def _read_file_preview(self, path: str, ftype: str, max_chars: int = 4000) -> str:
        try:
            if ftype == 'text':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    data = f.read(max_chars + 1024)  # 多讀少量以便截斷提示
                return self._truncate(data, max_chars) or "(空白)"

            if ftype == 'docx':
                if Document is None:
                    return "[缺少 python-docx 套件：pip install python-docx]"
                doc = Document(path)
                parts, length = [], 0
                for p in doc.paragraphs:
                    t = (p.text or '').strip()
                    if t:
                        parts.append(t)
                        length += len(t)
                    if length > max_chars:
                        break
                return self._truncate("\n".join(parts), max_chars) or "(空白 DOCX)"

            if ftype == 'doc':
                # 嘗試透過 Word COM（需要 Windows + 已安裝 Word + pywin32）
                if win32com is None:
                    return "[DOC 讀取需要 pywin32 以及安裝 Microsoft Word]"
                try:
                    word = win32com.gencache.EnsureDispatch('Word.Application')  # type: ignore[attr-defined]
                    word.Visible = False
                    doc = word.Documents.Open(path, ReadOnly=True)
                    text = doc.Content.Text if doc and doc.Content else ''
                    doc.Close(False)
                    word.Quit()
                    return self._truncate(text.replace("\r", "\n"), max_chars) or "(空白 DOC)"
                except Exception as e:
                    try:
                        # 盡力關閉殘留的 Word 實例
                        word.Quit()
                    except Exception:
                        pass
                    return f"[DOC 讀取失敗] {e}"

            if ftype == 'pdf':
                if PdfReader is None:
                    return "[缺少 pypdf 套件：pip install pypdf]"
                try:
                    reader = PdfReader(path)
                    parts, length = [], 0
                    for page in reader.pages:
                        try:
                            t = page.extract_text() or ''
                        except Exception:
                            t = ''
                        if t:
                            parts.append(t.strip())
                            length += len(t)
                        if length > max_chars:
                            break
                    return self._truncate("\n".join(parts), max_chars) or ""
                except Exception as e:
                    return f"[PDF 讀取錯誤] {e}"

            # 其他類型
            size = os.path.getsize(path)
            return f"[非文字類型，大小 {size} bytes]"
        except Exception as e:
            return f"[讀取錯誤] {e}"


    def _truncate(self, data: str, max_chars: int) -> str:
        if data is None:
            return ''
        if len(data) > max_chars:
            return data[:max_chars] + "\n...(截斷)"
        return data

    def _detect_syntax_type(self, ext: str) -> str:
        """根據副檔名偵測語法類型，用於語法高亮提示。"""
        syntax_map = {
            '.py': 'python',
            '.js': 'javascript', 
            '.ts': 'typescript',
            '.html': 'html',
            '.css': 'css',
            '.json': 'json',
            '.xml': 'xml',
            '.yml': 'yaml',
            '.yaml': 'yaml',
            '.md': 'markdown',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bat': 'batch',
            '.ps1': 'powershell',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.java': 'java',
            '.cs': 'csharp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust'
        }
        return syntax_map.get(ext, 'text')

    def _to_file_url(self, path: str) -> str:
        try:
            return Path(path).absolute().as_uri()
        except Exception:
            # 後備：手動組 file:///
            p = os.path.abspath(path).replace('\\', '/')
            if not p.startswith('/'):
                p = '/' + p
            return 'file://' + p

    def _render_pdf_pages(self, path: str, max_pages: int = 10, zoom: float = 1.5) -> list:
        """使用 PyMuPDF 產出多頁 PNG，回傳 file:// URL 陣列。無 PyMuPDF 則回空陣列。"""
        if fitz is None:
            return []
        try:
            doc = fitz.open(path)
            out_dir = tempfile.mkdtemp(prefix='AnoniMe_pdf_')
            urls = []
            page_count = min(len(doc), max_pages)
            mat = fitz.Matrix(zoom, zoom)
            for i in range(page_count):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_path = os.path.join(out_dir, f'page_{i+1}.png')
                pix.save(img_path)
                urls.append(Path(img_path).as_uri())
            doc.close()
            return urls
        except Exception:
            return []

    def _create_hex_view(self, path: str, base_data: dict, max_bytes: int = 1024) -> dict:
        """為二進位檔案創建十六進制檢視。"""
        try:
            with open(path, 'rb') as f:
                data = f.read(max_bytes)
            
            hex_lines = []
            for i in range(0, len(data), 16):
                chunk = data[i:i+16]
                hex_part = ' '.join(f'{b:02X}' for b in chunk)
                ascii_part = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
                hex_lines.append({
                    "offset": f"{i:08X}",
                    "hex": hex_part,
                    "ascii": ascii_part
                })
            
            return {
                **base_data,
                "viewType": "hex",
                "hexLines": hex_lines,
                "totalBytes": len(data),
                "isPartial": len(data) >= max_bytes
            }
        except Exception:
            return {**base_data, "viewType": "error", "error": "無法讀取檔案"}

    # 新增：建立預設的處理後輸出路徑（系統暫存目錄/AnoniMe/processed）
    def _make_processed_output_path(self, in_path: str) -> str:
        base = os.path.splitext(os.path.basename(in_path))[0]
        ext = os.path.splitext(in_path)[1].lower()
        if ext not in ('.pdf', '.docx', '.txt'):
            # 預設轉回原副檔名
            out_ext = ext or '.out'
        else:
            out_ext = ext
        out_dir = os.path.join(tempfile.gettempdir(), 'AnoniMe', 'processed')
        os.makedirs(out_dir, exist_ok=True)
        return os.path.join(out_dir, f"{base}_deid{out_ext}")

    # 新增：同步執行處理腳本，擷取 stdout 中的 file:// URL 並回傳 (url, path)
    def _run_processor_and_get_output(self, script_name: str, in_path: str, out_path: str | None = None) -> tuple[str | None, str | None]:
        script_path = Path(__file__).resolve().parent / script_name
        if not script_path.exists():
            print(f"[processor missing] {script_name} 不存在")
            return None, None
        try:
            cmd = [sys.executable, str(script_path), in_path]
            if out_path:
                cmd.append(out_path)
            res = subprocess.run(cmd, cwd=str(script_path.parent), capture_output=True, text=True, timeout=600)
            stdout = (res.stdout or '').strip()
            stderr = (res.stderr or '').strip()
            if res.returncode != 0:
                print(f"[processor nonzero] {script_name} rc={res.returncode}")
                if stderr:
                    print(stderr[:2000])
            # 1) 優先：找最後一行 file:// URL
            url = None
            if stdout:
                lines = [ln.strip() for ln in stdout.splitlines() if ln.strip()]
                for ln in reversed(lines):
                    if ln.lower().startswith('file://'):
                        url = ln
                        break
                # 2) 次要：行中包含 file:// 片段
                if url is None:
                    m = re.search(r'file://\S+', stdout, flags=re.IGNORECASE)
                    if m:
                        url = m.group(0)
                # 3) 次要：嘗試從 stdout 解析 Windows 路徑
                if url is None:
                    # 取最後一個存在的檔案路徑
                    for ln in reversed(lines):
                        # 移除常見前綴（例如：輸出檔案: ）
                        cand = re.sub(r'^[^A-Za-z0-9]+', '', ln)
                        # 簡單偵測 Windows/Unix 路徑
                        win_path = re.search(r'[A-Za-z]:\\[^\n\r]+', cand)
                        if win_path and os.path.exists(win_path.group(0)):
                            p = win_path.group(0)
                            try:
                                return Path(p).resolve().as_uri(), p
                            except Exception:
                                return None, None
                        if os.path.exists(cand):
                            try:
                                return Path(cand).resolve().as_uri(), cand
                            except Exception:
                                return None, None
            # 4) 若指定 out_path 且檔案存在，直接回傳
            if url is None and out_path and os.path.exists(out_path):
                try:
                    return Path(out_path).resolve().as_uri(), out_path
                except Exception:
                    pass
            if not url:
                print('[processor parse] 未取得輸出 URL，stdout/stderr:')
                if stdout:
                    print(stdout[:2000])
                if stderr:
                    print(stderr[:2000])
                return None, None
            # 轉成本機路徑
            path = self._file_url_to_path(url)
            if path is None:
                # 如果不是 URL，視為本機路徑並回推 URL
                path = url
                try:
                    url = Path(path).resolve().as_uri()
                except Exception:
                    url = None
            return url, path
        except Exception as e:
            print(f"[processor exec error] {script_name}: {e}")
            return None, None

    # 新增：file:// URL 轉成本機路徑（含 Windows 相容）
    def _file_url_to_path(self, url: str) -> str | None:
        try:
            if not url.lower().startswith('file:'):
                return None
            u = urlparse(url)
            p = unquote(u.path)
            if sys.platform == 'win32':
                # 處理 /C:/ 形式
                if p.startswith('/') and len(p) > 3 and p[2] == ':':
                    p = p[1:]
                return p.replace('/', '\\')
            return p
        except Exception:
            return None

    def _convert_office_to_pdf(self, path: str) -> str | None:
        """使用 Word COM 將 DOC/DOCX 轉為 PDF。僅在 Windows + 安裝 Word + pywin32 時可行。"""
        if sys.platform != 'win32' or win32com is None:
            return None
        word = None
        try:
            word = win32com.client.Dispatch('Word.Application')  # type: ignore[attr-defined]
            word.Visible = False
            doc = word.Documents.Open(path)
            # 將 PDF 輸出到工作目錄的 tmp_preview
            out_dir = os.path.join(os.getcwd(), 'tmp_preview')
            os.makedirs(out_dir, exist_ok=True)
            out_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(path))[0] + '.pdf')
            wdFormatPDF = 17
            doc.SaveAs(out_pdf, FileFormat=wdFormatPDF)
            doc.Close(False)
            return out_pdf
        except Exception:
            return None
        finally:
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

    @Slot(result=str)
    def getTempOutputDir(self) -> str:
        """回傳應用程式臨時輸出資料夾（供 QML 顯示或開啟）。"""
        try:
            out_dir = os.path.join(tempfile.gettempdir(), 'AnoniMe')
            os.makedirs(out_dir, exist_ok=True)
            return out_dir
        except Exception as e:
            return f"[tempdir error] {e}"

    @Slot()
    def createTempTestFile(self):
        """建立一個臨時測試檔並把結果透過 resultsReady 發到 ResultPage。"""
        try:
            out_dir = self.getTempOutputDir()
            fn = f"preview_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            p = os.path.join(out_dir, fn)
            with open(p, 'w', encoding='utf-8') as f:
                f.write(f"Temporary file created: {p}\n")
            payload = [{
                "fileName": os.path.basename(p),
                "type": "text",
                "originalText": f"Temporary file created: {p}",
                "maskedText": f"Temporary file created: {p}",
                "fileUrl": Path(p).as_uri()
            }]
            self.resultsReady.emit(json.dumps(payload, ensure_ascii=False))
        except Exception as e:
            payload = [{
                "fileName": "temp_error.txt",
                "type": "text",
                "originalText": f"[error] {e}",
                "maskedText": f"[error] {e}"
            }]
            self.resultsReady.emit(json.dumps(payload, ensure_ascii=False))

    @Slot(str, result=str)
    def readFileContent(self, file_path: str) -> str:
        """讀取檔案內容，供 ResultPage 使用"""
        try:
            # 如果是 file:// URL，先轉換為本地路徑
            if file_path.startswith('file://'):
                file_path = self._file_url_to_path(file_path) or file_path
            
            if not os.path.isfile(file_path):
                return f"[檔案不存在] {file_path}"
            
            # 根據副檔名判斷檔案類型
            ext = Path(file_path).suffix.lower()
            
            if ext == ".pdf":
                # PDF 檔案使用 pypdf 提取文字
                if PdfReader is None:
                    return "[PDF 讀取需要 pypdf 套件]"
                try:
                    reader = PdfReader(file_path)
                    text_parts = []
                    for page in reader.pages:
                        text_parts.append(page.extract_text() or "")
                    return "\n".join(text_parts) or "[PDF 無法提取文字內容]"
                except Exception as e:
                    return f"[PDF 讀取錯誤] {e}"
                    
            elif ext == ".docx":
                # DOCX 檔案使用 python-docx 提取文字
                if Document is None:
                    return "[DOCX 讀取需要 python-docx 套件]"
                try:
                    doc = Document(file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    return "\n".join(text_parts) or "[DOCX 無文字內容]"
                except Exception as e:
                    return f"[DOCX 讀取錯誤] {e}"
                    
            else:
                # 文字檔案直接讀取，保留原始格式
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    return content
                except UnicodeDecodeError:
                    # 如果 UTF-8 讀取失敗，嘗試其他編碼
                    try:
                        with open(file_path, 'r', encoding='big5', errors='ignore') as f:
                            content = f.read()
                        return content
                    except:
                        # 如果仍然失敗，可能是二進制檔案，返回錯誤消息
                        return "[檔案格式不支援直接檢視]"
        except Exception as e:
            return f"[讀取錯誤] {e}"

    def _file_url_to_path(self, url: str) -> str | None:
        """將 file:// URL 轉換為本地檔案路徑"""
        try:
            if not url.lower().startswith('file:'):
                return None
            u = urlparse(url)
            p = unquote(u.path)
            if sys.platform == 'win32':
                # 處理 /C:/ 形式
                if p.startswith('/') and len(p) > 3 and p[2] == ':':
                    p = p[1:]
                return p.replace('/', '\\')
            return p
        except Exception:
            return None
