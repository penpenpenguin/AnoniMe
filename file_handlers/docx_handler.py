# file_handlers/docx_handler.py

import os
from docx import Document
from pii_models.presidio_detector import detect_pii
from faker_models.presidio_replacer_plus import replace_pii
# from faker_models.ai_replacer import replace_entities
from faker_models.muiltAI_pii_replace import replace_entities, MappingStore, KuwaChatClient

class DocxHandler:
    """
    處理 .docx 檔案 in-place 去識別化，
    保留所有段落格式、run 樣式與表格結構。
    """
    # 新增：初始化 LlamaChatClient 和 MappingStore
    def __init__(self):                                   
        self.client = KuwaChatClient()
        self.mapping = MappingStore()

    def deidentify(self, input_path: str, output_path: str, selected_types: list[str] = None) -> str:
        """
        1. 讀取 input_path 的 Word 文件
        2. 針對每個 run.text 呼叫 detect_pii() + replace_pii()
        3. 處理表格內 cell
        4. 存成 output_path，並回傳它
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"找不到輸入檔：{input_path}")

        doc = Document(input_path)
        
        # 處理段落中的 runs
        for para in doc.paragraphs:
            # 合併整個段落的文字
            full_text = para.text
            if not full_text.strip():
                continue
            # 對完整段落進行實體偵測
            entities = detect_pii(full_text, language="en", score_threshold=0.6)
            print(f"處理段落：'{full_text}'", "\n")
            
            if entities:
                # 生成替換後的完整文字
                # new_full_text = replace_pii(full_text, entities)
                new_full_text = replace_entities(full_text, entities)
                new_full_text = replace_entities(                   # ★ (新)
                    full_text,
                    entities,
                    chat_client=self.client,
                    mapping=self.mapping,        # ★ 同一份 mapping，保持一致性
                    # batch_size=30,             # 可調；大量文件時 20~50 都可
                )
                print("替換後內容：", new_full_text)

                
                if new_full_text != full_text:
                    # 清空所有 runs 並重新設置文字
                    for run in para.runs:
                        run.text = ""
                    
                    # 將新文字設置到第一個 run（如果存在）
                    if para.runs:
                        para.runs[0].text = new_full_text
                    else:
                        # 如果沒有 runs，創建一個新的
                        para.add_run(new_full_text)

        # 處理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # 合併整個段落的文字
                        full_text = para.text
                        if not full_text.strip():
                            continue
                        # 對完整段落進行實體偵測
                        entities = detect_pii(full_text, language="en", score_threshold=0.6)
                        if entities:
                            # 生成替換後的完整文字
                            # new_full_text = replace_pii(full_text, entities)
                            new_full_text = replace_entities(full_text, entities)
                            new_full_text = replace_entities(                          # ★ (新)
                                full_text,
                                entities,
                                chat_client=self.client,
                                mapping=self.mapping,   # ★ 同一份映射
                            )
                            
                            if new_full_text != full_text:
                                # 清空所有 runs 並重新設置文字
                                for run in para.runs:
                                    run.text = ""
                                
                                # 將新文字設置到第一個 run
                                if para.runs:
                                    para.runs[0].text = new_full_text
                                else:
                                    para.add_run(new_full_text)

        # 儲存結果
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"儲存結果：{output_path}")
        return output_path
